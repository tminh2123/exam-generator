from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import re
import random

def parse(docx_path):
    doc = Document(docx_path)
    questions = []
    current = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        tags = re.findall(r"\[([^\]]+)\]", text)

        if len(tags) >= 4:
            if current:
                questions.append(current)
            current = {
                "tags": {
                    "chủ_đề": tags[1],  
                    "mức_độ": tags[2],
                    "mã": tags[3]
                },
                "elements": []
            }
            continue

        if not current:
            continue

        if text.startswith("[IMAGE:"):
            image_path = re.search(r"\[IMAGE:\s*(.*?)\]", text).group(1).strip()
            current["elements"].append({"type": "image", "path": image_path})
        else:
            runs_data = []
            for run in para.runs:
                runs_data.append({
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_name": run.font.name,
                    "font_size": run.font.size.pt if run.font.size else None,
                    "color": run.font.color.rgb
                })

            para_format = {
                "alignment": para.alignment,
                "line_spacing": para.paragraph_format.line_spacing,
                "space_before": para.paragraph_format.space_before,
                "space_after": para.paragraph_format.space_after
            }

            current["elements"].append({
                "type": "paragraph",
                "runs": runs_data,
                "para_format": para_format
            })

    if current:
        questions.append(current)

    print(f"Hoàn tất đọc. Tổng số câu hỏi: {len(questions)}")
    return questions


def filter_questions(questions, matrix):
    selected = []
    used_ids = set()

    for condition in matrix:
        filtered = []
        for q in questions:
            if id(q) in used_ids:
                continue
            match = True
            for key, value in condition.items():
                if key == "số_câu":
                    continue 
                if q["tags"].get(key) != value:
                    match = False
                    break
            if match:
                filtered.append(q)
        if len(filtered) < condition["số_câu"]:
            raise ValueError(f"❌ Không đủ câu hỏi cho {condition}")

        sampled = random.sample(filtered, condition["số_câu"])
        selected.extend(sampled)
        used_ids.update(id(q) for q in sampled)

    return selected

def save_exam(selected_questions, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(13)

    doc.add_heading("ĐỀ KIỂM TRA HÓA HỌC", 0)

    for i, q in enumerate(selected_questions, 1):
        para = doc.add_paragraph()
        fmt = para.paragraph_format
        fmt.line_spacing = 1      
        fmt.space_before = Pt(0)     
        fmt.space_after = Pt(0)      
        run = para.add_run(f"Câu {i}: ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True  

        for elem in q["elements"]:
            if elem["type"] == "paragraph":
                para = doc.add_paragraph()

                # Copy paragraph format
                if "para_format" in elem:
                    pf = elem["para_format"]
                    para.paragraph_format.alignment = pf.get("alignment")
                    para.paragraph_format.line_spacing = pf.get("line_spacing")
                    para.paragraph_format.space_before = pf.get("space_before")
                    para.paragraph_format.space_after = pf.get("space_after")

                # Copy runs
                for run_data in elem["runs"]:
                    run = para.add_run(run_data["text"])
                    run.bold = run_data["bold"]
                    run.italic = run_data["italic"]
                    run.underline = run_data["underline"]

                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

                    if run_data.get("font_size"):
                        run.font.size = Pt(run_data["font_size"])
                    if run_data.get("color"):
                        run.font.color.rgb = run_data["color"]

            elif elem["type"] == "image":
                try:
                    doc.add_picture(elem["path"], width=Inches(4))
                except Exception as e:
                    doc.add_paragraph(f"[LỖI CHÈN ẢNH: {elem['path']}] ({e})")

    doc.save(output_path)
    print(f"Đã lưu đề tại: {output_path}")


if __name__ == "__main__":
    all_questions = parse("bank.docx")

    ma_tran = [
        {"chủ_đề": "Polymer", "mức_độ": "B", "số_câu": 4},
        {"chủ_đề": "Polymer", "mức_độ": "H", "số_câu": 2},
        {"chủ_đề": "Polymer", "mức_độ": "VD", "số_câu": 2}
    ]

    selected = filter_questions(all_questions, ma_tran)
    save_exam(selected, "test.docx")
